import os
from time import sleep
import unicodedata
import pyperclip
import pyautogui, sys
from docx import Document
import json


def iterate_banks():
    docx_path = 'C:\\Users\\john\\Desktop\\banks_working_data\\'
    pyautogui.FAILSAFE = False
    
    json_directory = 'togo'
    json_files = os.listdir(json_directory)

    needs_scraped = []
    for json_file in json_files:
        if json_file == '.DS_Store':
            pass
        else:
            needs_scraped.append(json_file)

    print(len(needs_scraped))
    print(needs_scraped)
    for json_file in needs_scraped:
        with open('already_scraped.txt', 'a') as fout:
            fout.write(json_file + '\n')
 
    
        file_path = os.path.join(json_directory, json_file)
        with open(file_path, 'r') as fin:
            text = fin.read()
        
        text_components = text.split(',')
        #print(text_components)
        bank_name = text_components[0].split(':')[1].replace('"', '').strip()
        bank_text = text_components[1].split(':')[1].replace('"', '').replace('\\n', '. ').replace('..', '.').strip()
        #bank_num_pages = text_components[2].split(':')[1].replace('}', '').strip()

        #print(bank_name)
        #print(bank_text)
        #print(bank_num_pages)
        if len(bank_text) < 10:
            pass
        else:
            bank_file_name = bank_name.replace(' ', '_') + '.docx'
            document = Document()
            document.add_paragraph(bank_text)
            document.save(bank_file_name)

            document_length = bank_text.split()
            sleep_timer = (len(document_length) / 1000) * 10
            print(sleep_timer)
            if sleep_timer > 6500:
                with open('too_big.txt', 'a') as fout:
                    fout.write(json_file + '\n')
                os.remove(docx_path + bank_file_name)

                text_chunks = [document_length[x:x+350000] for x in range(0, len(document_length), 350000)]
                for i in range(len(text_chunks)):
                    print('Working on chunk ' + str(i))
                    chunk_text = ' '.join(text_chunks[i])
                    sleep_timer = (len(text_chunks[i]) / 1000) * 10
                    chunk_document_name = str(i) + '_' + bank_file_name
                    document = Document()
                    document.add_paragraph(chunk_text)
                    document.save(chunk_document_name)
                        
                    os.startfile(docx_path + chunk_document_name)
                    sleep(1.2)
                    pyautogui.click(849, 63)
                    sleep(1)
                    pyautogui.press(['down', 'enter'])
                    sleep(sleep_timer)
                    sleep(15)
                    pyautogui.click(884, 65)
                    sleep(1)
                    pyautogui.click(811, 16)
                    sleep(.5)
                    pyautogui.press(['right', 'enter'])
                    sleep(1)
                    os.remove(docx_path + chunk_document_name)
                    sleep(1)

                sleep(1)
                print()
            
            else:
                os.startfile(docx_path + bank_file_name)
                sleep(1.2)
                pyautogui.click(849, 63)
                sleep(1)
                pyautogui.press(['down', 'enter'])
                sleep(sleep_timer)
                sleep(15)
                pyautogui.click(884, 65)
                sleep(1)
                pyautogui.click(811, 16)
                sleep(.5)
                pyautogui.press(['right', 'enter'])
                sleep(1)
                os.remove(docx_path + bank_file_name)
                sleep(1)



def get_spot():
    try:
        while True:
            x, y = pyautogui.position()
            positionStr = 'X: ' + str(x).rjust(4) + ' Y: ' + str(y).rjust(4)
            print(positionStr, end='')
            print('\b' * len(positionStr), end='', flush=True)
    except KeyboardInterrupt:
        print('\n')


def count_bogged():
    scraped = []
    with open('sw4stats.txt', 'r') as fin:
        lines = fin.readlines()
        for line in lines:
            split_line = line.split('	')
            scraped.append(split_line[1].split('	')[0].replace('.docx', '.json'))
    with open('already_scraped.txt', 'w') as fout:
        for line in scraped:
            fout.write(line + '\n')


iterate_banks()
#get_spot()
#load_jsons()
#count_bogged()

